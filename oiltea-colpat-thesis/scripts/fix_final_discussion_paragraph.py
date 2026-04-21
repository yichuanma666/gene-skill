from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"F:\jyfx")
OUT_DIR = ROOT / "analysis_outputs"


def latest_final() -> Path:
    candidates = sorted(
        [p for p in OUT_DIR.glob("*.docx") if p.name.endswith("_1.3w_no_fig_titles.docx")],
        key=lambda x: x.stat().st_mtime,
    )
    if candidates:
        return candidates[-1]
    return sorted(
        [p for p in OUT_DIR.glob("*.docx") if not p.name.startswith("CoLPAT_")],
        key=lambda x: x.stat().st_mtime,
    )[-1]


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def set_run_font(run, size=12, bold=False, chinese="宋体", western="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = western
    rpr = run._element.get_or_add_rPr()
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), chinese)
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    rpr.append(fonts)


def format_body(paragraph):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.first_line_indent = Pt(24)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def main():
    path = latest_final()
    doc = Document(path)
    replacement = (
        "此外，本研究在论文写作中将表达结果明确限定为公开群体材料证据，避免将其误写为本课题组织样品或发育时期表达结果，这一点对于保证结果真实性较为关键。"
        "后续若获得普通油茶根、茎、叶、花及不同种子发育时期的 TPM/FPKM 矩阵，可在当前成员鉴定和结构域验证结果基础上直接替换表达热图，并进一步结合含油率、脂肪酸组分和候选基因表达量进行相关性分析，从而把本研究由生物信息学鉴定推进到功能验证和育种应用层面。"
        "同时，当前形成的 CoLPAT 成员表、PF01553 结构域坐标和公开表达映射结果也可作为后续实验设计的基础资料，用于引物设计、候选成员优先级排序以及不同材料间表达差异比较。"
        "这些结果具有较好的连续利用价值。因此，本文结果既能支撑毕业论文的系统分析，也能为后续课题继续深化提供可复用的数据基础。"
    )
    target = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "TPM/FPKM" in text and (text.count("?") > 20 or text.startswith("此外，本研究在论文写作中")):
            target = paragraph
            break
    if target is None:
        raise RuntimeError("Target discussion paragraph not found")
    clear_paragraph(target)
    run = target.add_run(replacement)
    set_run_font(run)
    format_body(target)
    doc.save(path)
    chars = len("\n".join(p.text for p in doc.paragraphs).replace(" ", "").replace("\n", ""))
    print(path)
    print(chars)


if __name__ == "__main__":
    main()
