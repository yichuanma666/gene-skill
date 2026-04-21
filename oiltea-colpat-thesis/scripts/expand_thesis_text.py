from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"F:\jyfx")
OUT_DIR = ROOT / "analysis_outputs"


def latest_thesis() -> Path:
    candidates = sorted(
        [p for p in OUT_DIR.glob("*.docx") if not p.name.startswith("CoLPAT_")],
        key=lambda x: x.stat().st_mtime,
    )
    if not candidates:
        raise FileNotFoundError("No thesis docx found")
    return candidates[-1]


def paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def set_run_font(run, size=12, bold=False, chinese="宋体", western="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = western
    rpr = run._element.get_or_add_rPr()
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), chinese)
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    rpr.append(fonts)


def format_body(paragraph):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.first_line_indent = Pt(24)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_body_after(doc: Document, index: int, text: str, style):
    paragraph = doc.paragraphs[index]
    nxt = paragraph._element.getnext()
    if nxt is not None:
        nxt_para = Paragraph(nxt, paragraph._parent)
        next_text = nxt_para.text or ""
        if next_text.strip().startswith(text[:18]):
            return
    new_para = paragraph_after(paragraph)
    new_para.style = style
    run = new_para.add_run(text)
    set_run_font(run, size=12)
    format_body(new_para)


def main():
    path = latest_thesis()
    doc = Document(path)
    body_style = doc.paragraphs[37].style

    additions = [
        (
            37,
            "从产业发展需求看，普通油茶优质高产新品种的选育不仅取决于产量性状，还受到含油量、脂肪酸组成、逆境适应能力和稳产性的共同影响。随着油茶高值化利用和全产业链开发持续推进，围绕油脂合成关键酶开展分子水平解析，已经成为连接基因资源挖掘与分子育种应用的重要环节。相比传统性状调查，基因家族层面的系统鉴定更有助于从全基因组尺度筛选候选功能基因，建立后续表达验证、关联分析和功能转化的基础框架。",
        ),
        (
            39,
            "已有研究表明，不同植物中的 LPAT 成员在亚细胞定位、底物偏好和表达模式上常表现出明显分工。例如，部分成员偏向参与内质网途径中的储油脂合成，部分成员则与叶绿体膜脂组装、胚乳发育或逆境响应有关。因此，在油茶中开展 LPAT 家族的系统比较，不仅能够明确成员数量和保守结构特征，还能为推测不同成员在茶油品质形成中的潜在作用提供依据。",
        ),
        (
            41,
            "与仅进行家族鉴定不同，本研究同时引入公开表达矩阵和拟南芥种间共线性信息，对 CoLPAT 成员的结构、进化和表达特征进行联动分析，力求在现有数据条件下尽可能提高结果的解释深度。这种分析框架既可服务于当前论文写作，也为后续补充真实组织表达谱、脂肪酸测定数据和功能验证实验预留了接口。",
        ),
        (
            61,
            "从染色体分布看，16 个 CoLPAT 成员并非集中成簇分布，而是分散定位于不同染色体或单倍型伪染色体上，仅在部分同源染色体组中表现出相对对应的保留关系。这种分布格局说明，普通油茶 LPAT 家族在多倍化后并未发生大规模局部串联扩增，而更可能是在基因组复制与后续重排过程中保留了部分具有功能价值的拷贝。不同成员在氨基酸长度和理化性质上的差异，也提示该家族在长期进化中逐步形成了功能储备。",
        ),
        (
            68,
            "系统树中普通油茶成员主要聚集在 AtLPAT1、AtLPAT2 和 AtLPAT4 相关分支，而未形成与 AtLPAT3、AtLPAT5 数量相当的独立扩张簇，说明 CoLPAT 家族的扩张具有明显偏向性。考虑到不同 AtLPAT 分支在膜脂代谢和甘油三酯合成中的功能侧重不同，这种聚类结果提示普通油茶中与种子油脂合成密切相关的 LPAT 功能模块可能主要集中在 AtLPAT1/2/4 同源成员之中，值得后续优先验证。",
        ),
        (
            72,
            "结构域复核结果进一步提高了成员鉴定的可靠性。PF01553 是当前 CoLPAT 鉴定中最核心的 HMM 证据，而 CDD 与 SMART 的交叉验证则有助于排除仅靠低复杂度区段或局部短片段误判的候选序列。对于家族鉴定类研究而言，这种“模型筛选 + 结构域复核”的流程能够有效保证最终成员集合的稳定性，也使后续系统发育、结构比较和表达分析建立在更可信的成员基础之上。",
        ),
        (
            79,
            "同源类别内成员在外显子组织形式上的相似性，与系统发育树中的近缘聚类关系相互印证，说明普通油茶 LPAT 家族在进化过程中保留了较为稳定的结构骨架。另一方面，个别成员在基因长度和内含子数量上明显偏离家族平均水平，这类结构差异常常与转录调控复杂化、可变剪接潜力或表达专化有关，值得在后续转录本水平分析中进一步关注。",
        ),
        (
            84,
            "需要指出的是，最佳同源关系和严格共线性关系并不完全等价。前者更多反映蛋白序列层面的近缘性，后者则要求保留较完整的基因邻域顺序信息。普通油茶中多数 CoLPAT 成员虽然能够找到相对稳定的拟南芥同源分支，但只有少数位点保留严格 WGDI 共线性证据，说明该家族在种间比较中既具有保守的一面，也受到了多倍化和基因组重塑的明显影响。",
        ),
        (
            94,
            "从表达层次看，CoLPAT 成员之间并非简单的高低表达关系，而更可能存在不同材料背景下的差异化调控模式。部分成员在高变异材料中呈现相似的升降趋势，提示其可能参与同一代谢模块；另一些成员则整体表达较低，可能承担更为专一或特定条件下才被激活的功能。公开表达结果虽然不能直接替代组织时空表达谱，但已经为候选基因优先级排序提供了较有价值的依据。",
        ),
        (
            99,
            "与年生模式植物相比，油茶属于多年生木本油料作物，其长期营养生长、开花结实和种子充实过程更加复杂，这也意味着脂质代谢调控网络可能具有更强的阶段性和组织特异性。因此，在油茶中系统鉴定 LPAT 家族，不仅是简单补充一个物种的家族成员列表，更是在木本油料作物背景下重新审视脂质合成关键酶进化与功能分工的重要基础工作。",
        ),
        (
            100,
            "从方法学角度看，本文对 CoLPAT 的最终认定并未仅依赖单一同源比对结果，而是综合了 HMM 模型、理化性质、结构域复核和系统发育关系等多重证据。这种多证据交叉支持的策略，使得最终保留下来的 16 个成员在结构和进化上都具有较好的自洽性，也提高了结论在论文写作和后续功能验证中的可引用性。",
        ),
        (
            101,
            "对于候选基因筛选而言，兼顾进化保守性和表达活性通常比单看某一指标更有参考价值。严格共线性支持的 CoLPAT16 更适合用于探讨家族祖先保守功能，而在公开群体中表达较高的 CoLPAT1、CoLPAT2、CoLPAT4、CoLPAT6、CoLPAT7 和 CoLPAT8，则更适合作为优先开展组织表达检测、相关性分析和酶学验证的对象。将两类证据结合起来，有助于缩小后续实验验证范围，提高研究效率。",
        ),
    ]

    for index, text in sorted(additions, key=lambda item: item[0], reverse=True):
        add_body_after(doc, index, text, body_style)

    doc.save(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    chars = len(text.replace(" ", "").replace("\n", ""))
    print(path)
    print(chars)


if __name__ == "__main__":
    main()
