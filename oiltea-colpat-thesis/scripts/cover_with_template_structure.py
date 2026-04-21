from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"F:\jyfx")
OUT_DIR = ROOT / "analysis_outputs"


def latest_thesis() -> Path:
    candidates = sorted(
        [p for p in OUT_DIR.glob("*.docx") if not p.name.startswith("CoLPAT_")],
        key=lambda x: x.stat().st_mtime,
    )
    if not candidates:
        raise FileNotFoundError("No thesis docx found")
    return candidates[-1]


def paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def find_para(doc: Document, startswith: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise ValueError(f"Paragraph not found: {startswith}")


def has_drawing(paragraph) -> bool:
    return "w:drawing" in paragraph._element.xml


def remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    parent.remove(p)
    paragraph._p = paragraph._element = None


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def set_run_font(run, size=12, bold=False, chinese="宋体", western="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = western
    rpr = run._element.get_or_add_rPr()
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), chinese)
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    rpr.append(fonts)


def format_body(paragraph):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.first_line_indent = Pt(24)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def format_heading(paragraph):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.first_line_indent = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def format_caption(paragraph):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.first_line_indent = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_body_text(paragraph, text, style=None):
    if style is not None:
        paragraph.style = style
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)
    format_body(paragraph)


def set_heading_text(paragraph, text, style=None):
    if style is not None:
        paragraph.style = style
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=12, bold=False)
    format_heading(paragraph)


def set_caption_text(paragraph, text, style=None):
    if style is not None:
        paragraph.style = style
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    format_caption(paragraph)


def delete_figure_by_caption(doc: Document, caption_startswith: str):
    cap = find_para(doc, caption_startswith)
    prev = cap._element.getprevious()
    if prev is not None:
        prev_para = Paragraph(prev, cap._parent)
        if has_drawing(prev_para):
            remove_paragraph(prev_para)
    remove_paragraph(cap)


def main():
    path = latest_thesis()
    doc = Document(path)

    top_heading_style = find_para(doc, "1 引言").style
    sub_heading_style = find_para(doc, "2.1").style
    body_style = find_para(doc, "普通油茶（Camellia oleifera Abel.）属于山茶科山茶属").style
    caption_style = find_para(doc, "图1").style

    # Remove figures that do not match the reference-paper mainline.
    for caption in [
        "图10",
        "图8  普通油茶 CoLPAT 蛋白二级结构预测",
        "图6 拟南芥AtLPAT与普通油茶CoLPAT最佳同源关系示意图",
    ]:
        delete_figure_by_caption(doc, caption)

    # Intro rewritten to match the reference-paper section arrangement.
    p_species = find_para(doc, "普通油茶（Camellia oleifera Abel.）属于山茶科山茶属")
    t_species = p_species.text.strip()
    set_heading_text(p_species, "1.1 油茶的产业价值及脂质合成研究背景", sub_heading_style)
    p_species_body = paragraph_after(p_species)
    set_body_text(p_species_body, t_species, body_style)

    p_progress = find_para(doc, "三酰甘油（TAG）是油茶种子中主要的储存脂质")
    t_progress_1 = p_progress.text.strip()
    p_progress_body = find_para(doc, "目前，拟南芥、水稻、油菜、棉花等植物的LPAT家族已完成")
    t_progress_2 = p_progress_body.text.strip()
    set_heading_text(p_progress, "1.2 LPAT基因家族研究进展", sub_heading_style)
    set_body_text(
        p_progress_body,
        t_progress_1
        + " "
        + t_progress_2
        + " 作为 Kennedy 途径中的关键酰基转移酶，LPAT 的家族成员鉴定和表达差异分析，对于解析油茶种子油脂积累机制具有直接意义。",
        body_style,
    )

    p_goal = find_para(doc, "本研究以普通油茶“长林40号”基因组为基础，对 LPAT 基因家族成员进行系统鉴定")
    t_goal = p_goal.text.strip()
    set_heading_text(p_goal, "1.3 研究目的与意义", sub_heading_style)
    p_goal_body = paragraph_after(p_goal)
    set_body_text(p_goal_body, t_goal, body_style)

    # Methods.
    set_body_text(
        find_para(doc, "以普通油茶“长林40号”基因组注释文件为基础"),
        "普通油茶“长林40号”基因组蛋白序列、CDS 序列和 GFF3 注释文件用于 CoLPAT 家族鉴定。首先下载 Pfam 数据库中的 PF01553（Acyltransferase）隐马尔可夫模型，对全蛋白集进行 HMMER/pyhmmer 检索；同时以拟南芥 AtLPAT 蛋白序列为参考，利用 DIAMOND/BLASTP 进行同源筛选。对获得的候选序列去冗余后，再结合 NCBI Batch CD-Search、CDD 家族注释规则和 SMART 域注释对保守结构域进行复核，最终确定普通油茶 LPAT 基因家族成员，并按染色体位置依次命名为 CoLPAT1-CoLPAT16。",
        body_style,
    )
    set_body_text(
        find_para(doc, "将CoLPAT蛋白序列分别提交NCBI-CDD和SMART数据库"),
        "将 CoLPAT 蛋白序列提交 NCBI Batch CD-Search 和 SMART 数据库进行保守结构域验证，重点确认是否含有 LPAT/LPLAT 相关 Acyltransferase 保守结构域。由于当前分析流程以 PF01553 HMM 命中区段为核心，保守基序图采用 PF01553 命中坐标及相邻保守区块在 TBtools 中进行统一可视化，用于比较不同成员保守区块的数量、排列顺序和保守程度。",
        body_style,
    )
    set_body_text(
        find_para(doc, "下载拟南芥 TAIR10 基因组注释文件和蛋白序列"),
        "下载拟南芥 TAIR10 基因组注释文件和蛋白序列，提取 AtLPAT 基因位置信息。采用 DIAMOND 对普通油茶与拟南芥最长蛋白序列进行全蛋白同源比对，并利用 WGDI 识别两物种之间的共线性区块。将 CoLPAT 与 AtLPAT 的坐标映射到共线性结果中，筛选涉及 CoLPAT 或 AtLPAT 的相关区块，以评估普通油茶 LPAT 家族与拟南芥 LPAT 家族之间的保守同源关系。",
        body_style,
    )
    set_body_text(
        find_para(doc, "从 GEO 数据库下载普通油茶公开转录组项目 GSE190644 的 FPKM 表达矩阵及转录本序列"),
        "从 GEO 数据库下载普通油茶公开转录组项目 GSE190644 的 FPKM 表达矩阵及转录本序列。将转录本序列翻译为蛋白序列后，利用 PF01553 隐马尔可夫模型筛选 LPAT 同源转录本，并通过 DIAMOND 将 16 个 CoLPAT 蛋白映射到公共表达数据中的对应转录本。表达量以 log2(FPKM+1) 转换，并按基因进行行标准化后绘制表达热图，同时统计 221 份材料中的中位 FPKM 作为辅助表达指标。",
        body_style,
    )

    # Results.
    set_body_text(
        find_para(doc, "基于长林40号基因组序列和GFF3注释文件提取全基因组预测蛋白"),
        "采用拟南芥 AtLPAT 同源序列检索和 PF01553（Acyltransferase）隐马尔可夫模型双重筛选，并结合 CDD 与 SMART 对结构域进行复核，最终在普通油茶基因组中鉴定到 16 个 CoLPAT 基因家族成员。染色体定位结果如图1所示，成员基本信息见表1。16 条 CoLPAT 蛋白序列编码氨基酸数为 277-920 aa，分子质量为 31.87-100.33 kDa，理论等电点范围为 6.56-10.37，不稳定系数为 38.02-129.72，GRAVY 为 -0.692-0.198，表明不同成员在蛋白大小、稳定性和疏水性方面存在一定差异。",
        body_style,
    )
    set_body_text(
        find_para(doc, "保守结构域分析显示"),
        "为了更好地了解 CoLPAT 蛋白序列的结构特征，对 16 个成员进行了保守区块和保守结构域分析。结果如图3和图4所示，16 个 CoLPAT 蛋白均可被 PF01553（Acyltransferase）模型检出，HMM 检索 E-value 范围为 1.61e-25-6.83e-08。结合 CDD 和 SMART 的复核结果，这些蛋白均归属于 LPAT/LPLAT 相关酰基转移酶保守结构域类型，未见与 LPAT 主功能明显冲突的主导额外结构域。不同成员之间保守区块的数量和排列总体相似，而同一拟南芥同源类别中的成员保守性更高，提示 CoLPAT 家族在保留核心催化功能的同时发生了一定程度的分化。",
        body_style,
    )
    set_body_text(
        find_para(doc, "根据 AtLPAT 与 CoLPAT 最佳同源关系构建的连接示意图显示"),
        "为了进一步了解普通油茶 CoLPAT 家族与拟南芥 LPAT 家族之间的同源保守关系，以拟南芥为比较物种开展种间共线性分析。结果如图6所示，普通油茶与拟南芥全基因组范围内共检测到 1370 个共线性区块，其中与 CoLPAT 或 AtLPAT 位点相关的区块有 5 个。CoLPAT16 与 AtLPAT3 位于同一共线性区块，说明二者具有较强的保守性；CoLPAT2、CoLPAT4、CoLPAT6 和 CoLPAT7 也位于种间共线性区段，但对应拟南芥区段内未覆盖 AtLPAT 位点，提示普通油茶 LPAT 家族在保守性基础上经历了不同程度的基因重排与分化。",
        body_style,
    )
    set_body_text(
        find_para(doc, "CoLPAT蛋白二级结构预测结果显示"),
        "通过二级结构预测对 16 个 CoLPAT 蛋白进行分析，结果如表2所示。CoLPAT 蛋白二级结构主要由 α-螺旋、延伸链、β-转角和无规则卷曲四种结构元素组成，其中 α-螺旋比例为 16.63%-35.62%，延伸链比例为 27.28%-41.83%，β-转角比例为 22.02%-56.41%，无规则卷曲比例为 0.00%-10.94%。整体来看，CoLPAT 蛋白以规则结构元件为主，无规则卷曲占比较低，说明家族整体折叠框架较为稳定；不同成员在 α-螺旋、延伸链和 β-转角比例上的差异，则提示其局部空间构象可能存在一定分化。",
        body_style,
    )
    set_body_text(
        find_para(doc, "基于 GSE190644 公开表达矩阵的补充分析表明"),
        "为了进一步了解 CoLPAT 家族成员在公开群体材料中的表达差异，对 GSE190644 表达矩阵进行了表达模式分析。结果如图7所示，16 个 CoLPAT 成员均可映射到油茶公共转录本中的 LPAT 同源序列，其中 13 个为高置信映射，3 个为低置信映射。在表达变异较大的 40 份材料中，不同 CoLPAT 同源转录本的表达模式存在明显差异，但同一拟南芥同源类别成员之间仍表现出一定相似性。结合 221 份材料的中位 FPKM 统计，CoLPAT2、CoLPAT4 和 CoLPAT1 对应转录本整体表达水平较高，CoLPAT8、CoLPAT6 和 CoLPAT7 次之，而 CoLPAT13-CoLPAT16 整体表达偏低。需要说明的是，该结果基于公开群体表达矩阵，反映的是材料间总体表达趋势，不能直接替代特定组织或发育时期的表达验证。",
        body_style,
    )

    # Rename remaining captions to match the reference-paper order.
    set_caption_text(
        find_para(doc, "图7  普通油茶 CoLPAT 与拟南芥 AtLPAT 相关 WGDI 共线性块"),
        "图6 普通油茶（Camellia oleifera）和拟南芥（Arabidopsis thaliana）LPAT基因家族种间共线性分析",
        caption_style,
    )
    set_caption_text(
        find_para(doc, "表2 普通油茶LPAT家族成员蛋白二级结构组成统计"),
        "表2 普通油茶 CoLPAT 家族成员蛋白二级结构占比",
        caption_style,
    )
    set_caption_text(
        find_para(doc, "图9  GSE190644 中 CoLPAT 同源转录本在高变异材料中的表达热图"),
        "图7 普通油茶（Camellia oleifera）CoLPAT同源转录本表达模式分析",
        caption_style,
    )

    # Discussion and conclusion rewritten to match the reference-paper layout.
    p_discussion = find_para(doc, "4 讨论")
    set_heading_text(p_discussion, "4 讨论与结论", top_heading_style)
    p_discuss_sub = paragraph_after(p_discussion)
    set_heading_text(p_discuss_sub, "4.1 讨论", sub_heading_style)

    set_body_text(
        find_para(doc, "LPAT 是植物甘油脂合成途径中的关键酰基转移酶"),
        "LPAT 是 Kennedy 途径中的关键酰基转移酶，其家族成员数量、结构特征与表达分化直接影响植物油脂合成效率。本研究在普通油茶中鉴定到 16 个 CoLPAT 成员，数量高于拟南芥等模式植物，说明在多倍化背景下油茶 LPAT 家族发生了扩张并得到一定程度的保留，这与油茶复杂基因组历史具有一致性。",
        body_style,
    )
    set_body_text(
        find_para(doc, "系统发育、保守结构域和基因结构分析表明"),
        "系统发育、保守结构域和基因结构分析表明，普通油茶 LPAT 家族整体较为保守。16 个 CoLPAT 成员均稳定检出 PF01553 保守结构域，且经 CDD/SMART 复核后均支持其属于 LPAT/LPLAT 相关酰基转移酶类型。与此同时，不同成员在外显子数目、蛋白长度和二级结构比例上仍存在差异，提示其在保持核心催化功能的同时，可能在底物偏好、亚细胞环境适应或调控方式上发生了分化。",
        body_style,
    )
    set_body_text(
        find_para(doc, "公开表达数据进一步显示"),
        "种间共线性和公开表达结果进一步表明，CoLPAT 家族既保留了与拟南芥 LPAT 的共同祖先信息，也表现出明显的物种特异性分化。WGDI 分析仅在少数位点检测到严格共线性关系，其中 CoLPAT16 与 AtLPAT3 的保守性最强；表达分析则显示 CoLPAT1、CoLPAT2、CoLPAT4、CoLPAT6、CoLPAT7 和 CoLPAT8 在公开群体中整体表达较高，可作为后续开展组织表达验证和功能研究的重点候选对象。",
        body_style,
    )
    set_body_text(
        find_para(doc, "受限于当前可获得的数据类型"),
        "本研究的表达分析基于公开群体材料数据，尚不能替代普通油茶不同组织、不同发育时期的真实表达谱，因此后续仍需结合特定组织样品、种子发育时期 RNA-seq、qRT-PCR、亚细胞定位和遗传转化等实验，对 CoLPAT 候选成员的生物学功能进行进一步验证。",
        body_style,
    )

    p_conclusion = find_para(doc, "5 结论")
    set_heading_text(p_conclusion, "4.2 结论", sub_heading_style)
    p_conclusion_1 = find_para(doc, "本研究基于普通油茶“长林40号”基因组共鉴定出 16 个 CoLPAT 基因家族成员")
    set_body_text(
        p_conclusion_1,
        "1. 从普通油茶“长林40号”基因组中共鉴定出 16 个 CoLPAT 基因家族成员，分布于 12 条染色体或单倍型伪染色体上；系统进化结果表明，这些成员主要与拟南芥 AtLPAT1、AtLPAT2 和 AtLPAT4 三类分支关系较近。",
        body_style,
    )
    p_conclusion_2 = paragraph_after(p_conclusion_1)
    set_body_text(
        p_conclusion_2,
        "2. 理化性质分析发现，16 条 CoLPAT 蛋白序列编码氨基酸数为 277-920 aa，分子质量为 31.87-100.33 kDa，理论等电点范围为 6.56-10.37，不稳定系数为 38.02-129.72，表明不同成员在蛋白稳定性和理化特征上存在一定差异。",
        body_style,
    )
    p_conclusion_3 = paragraph_after(p_conclusion_2)
    set_body_text(
        p_conclusion_3,
        "3. 保守区块和结构域分析显示，16 个 CoLPAT 成员均含有 PF01553 保守结构域；结合 CDD 与 SMART 复核结果，均支持其属于 LPAT/LPLAT 相关酰基转移酶类型，说明 CoLPAT 家族在核心结构域层面具有较高保守性。",
        body_style,
    )
    p_conclusion_4 = paragraph_after(p_conclusion_3)
    set_body_text(
        p_conclusion_4,
        "4. CoLPAT 蛋白二级结构主要由 α-螺旋、延伸链、β-转角和无规则卷曲构成，整体以规则结构元件为主，表明该家族蛋白折叠框架相对稳定，但不同成员之间仍存在一定的结构比例差异。",
        body_style,
    )
    p_conclusion_5 = paragraph_after(p_conclusion_4)
    set_body_text(
        p_conclusion_5,
        "5. 公开表达和种间共线性分析表明，CoLPAT2、CoLPAT4、CoLPAT1、CoLPAT8、CoLPAT6 和 CoLPAT7 的同源转录本在群体材料中整体表达较高；普通油茶与拟南芥之间共检测到 5 个与 CoLPAT/AtLPAT 相关的共线性区块，其中 CoLPAT16 与 AtLPAT3 具有严格共线性关系，可为后续功能验证提供重点候选基因和进化线索。",
        body_style,
    )

    doc.save(path)
    print(path)


if __name__ == "__main__":
    main()
