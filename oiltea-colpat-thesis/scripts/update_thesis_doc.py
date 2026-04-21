from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"F:\jyfx")
OUT_DIR = ROOT / "analysis_outputs"
FIG_DIR = OUT_DIR / "figures_ref_style"


def find_source_doc() -> Path:
    candidates = sorted(
        p for p in OUT_DIR.glob("*.docx") if not p.name.startswith("CoLPAT_")
    )
    if not candidates:
        raise FileNotFoundError("No thesis docx found in analysis_outputs")
    return candidates[0]


def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def set_run_fonts(run, size=12, bold=False, chinese="宋体", western="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = western
    rpr = run._element.get_or_add_rPr()
    east_asia = OxmlElement("w:rFonts")
    east_asia.set(qn("w:eastAsia"), chinese)
    east_asia.set(qn("w:ascii"), western)
    east_asia.set(qn("w:hAnsi"), western)
    rpr.append(east_asia)


def qn(tag: str) -> str:
    prefix, tagroot = tag.split(":")
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return f"{{{nsmap[prefix]}}}{tagroot}"


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def format_body(paragraph, first_line=True, center=False):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.first_line_indent = Pt(24) if first_line else Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY


def format_caption(paragraph):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.first_line_indent = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_paragraph_text(paragraph, text, *, size=12, bold=False, first_line=True, center=False):
    style = paragraph.style
    clear_paragraph(paragraph)
    paragraph.style = style
    run = paragraph.add_run(text)
    set_run_fonts(run, size=size, bold=bold)
    format_body(paragraph, first_line=first_line, center=center)


def set_caption_text(paragraph, text):
    style = paragraph.style
    clear_paragraph(paragraph)
    paragraph.style = style
    run = paragraph.add_run(text)
    set_run_fonts(run, size=10.5)
    format_caption(paragraph)


def paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_picture_with_caption(after_paragraph, image_path: Path, caption: str, width_cm: float):
    pic_para = paragraph_after(after_paragraph)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.add_run().add_picture(str(image_path), width=Cm(width_cm))
    cap_para = paragraph_after(pic_para)
    set_caption_text(cap_para, caption)
    return cap_para


def remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    parent.remove(p)
    paragraph._p = paragraph._element = None


def delete_from_heading(doc: Document, heading_text: str):
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(heading_text):
            start = i
            break
    if start is None:
        return
    for p in list(doc.paragraphs[start:]):
        remove_paragraph(p)


def find_paragraph_index(doc: Document, startswith: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(startswith):
            return i
    raise ValueError(f"Paragraph starting with {startswith!r} not found")


def main():
    src = find_source_doc()
    doc = Document(src)

    replacements = {
        23: "油茶（Camellia oleifera Abel.）是我国重要木本食用油树种，解析其脂质合成相关基因家族对于阐明油脂积累分子基础具有重要意义。溶血磷脂酸酰基转移酶（lysophosphatidic acid acyltransferase，LPAT）是 Kennedy 途径中催化磷脂酸形成的关键酶之一。本研究以普通油茶“长林40号”基因组为基础，对 LPAT 基因家族进行了全基因组鉴定、染色体定位、系统发育、保守结构域、基因结构及二级结构分析，并结合 GSE190644 公开表达矩阵与拟南芥 TAIR10 基因组数据补充分析其表达与种间共线性特征。结果表明：普通油茶基因组中共鉴定出 16 个 CoLPAT 成员，不均匀分布于 12 条染色体/单倍型伪染色体上；所有成员均含有 PF01553 保守结构域，系统发育上分别与 AtLPAT1、AtLPAT2 和 AtLPAT4 亲缘关系较近；基因结构和蛋白二级结构总体保守，但不同成员在内含子数目和螺旋/无规卷曲比例上存在差异。公开表达映射显示，16 个 CoLPAT 成员均可对应到油茶公共转录本中的 LPAT 同源序列，其中 13 个为高置信映射，CoLPAT1、CoLPAT2、CoLPAT4、CoLPAT6、CoLPAT7 和 CoLPAT8 在 221 份材料中表现出较高的中位 FPKM。WGDI 共检测到 5 个涉及 CoLPAT 或 AtLPAT 的相关共线性块，其中 CoLPAT16 与 AtLPAT3 位于同一共线性块。研究结果为普通油茶 LPAT 家族功能解析及油脂品质改良候选基因筛选提供了基础资料。",
        28: "Camellia oleifera is an important woody oil crop in China, and identifying lipid biosynthesis-related gene families is essential for understanding the molecular basis of oil accumulation. Lysophosphatidic acid acyltransferase (LPAT) is a key enzyme in the Kennedy pathway that catalyzes phosphatidic acid formation. In this study, the LPAT family was identified genome-wide from the Camellia oleifera cultivar Changlin40 genome, followed by chromosome localization, phylogenetic analysis, conserved domain analysis, gene structure analysis, and secondary structure prediction. Public expression data from GSE190644 and the Arabidopsis thaliana TAIR10 genome were further used to assess expression divergence and interspecific synteny. A total of 16 CoLPAT genes were identified and unevenly distributed on 12 chromosomes or haplotype pseudochromosomes. All members contained the conserved PF01553 acyltransferase domain and were mainly clustered with AtLPAT1-, AtLPAT2-, or AtLPAT4-related lineages. Gene structures and predicted secondary structures were generally conserved, although variation existed in intron numbers and the proportions of alpha-helix and random coil. Public transcriptome mapping indicated that all 16 CoLPAT members could be linked to LPAT-homologous transcripts in oil tea, including 13 high-confidence mappings, and CoLPAT1, CoLPAT2, CoLPAT4, CoLPAT6, CoLPAT7, and CoLPAT8 showed relatively high median FPKM values across 221 accessions. WGDI detected five synteny blocks involving CoLPAT or AtLPAT loci, among which CoLPAT16 and AtLPAT3 were located in the same block. These results provide a basis for functional characterization of the CoLPAT family and candidate gene selection for oil-quality improvement in C. oleifera.",
        39: "本研究以普通油茶“长林40号”基因组为基础，对 LPAT 基因家族成员进行系统鉴定，并从染色体分布、系统发育关系、保守结构域、基因结构、蛋白二级结构、公开表达模式及与拟南芥的共线性关系等方面开展综合分析，旨在为普通油茶脂质合成调控机制研究和分子育种提供候选基因与理论基础。",
        42: "以普通油茶“长林40号”基因组注释文件为基础，提取蛋白序列、CDS 序列和 GFF3 注释信息。下载 Pfam 数据库中的 PF01553（Acyltransferase）隐马尔可夫模型，采用 HMMER/pyhmmer 对普通油茶蛋白集进行筛选；同时以拟南芥 LPAT 蛋白序列为参考，利用 DIAMOND/BLASTP 进行同源检索。整合结构域筛选结果与同源比对结果，去除冗余序列后，再结合 NCBI CDD、SMART 等数据库对候选蛋白保守结构域进行复核，最终确定普通油茶 LPAT 基因家族成员，并按染色体位置顺序命名为 CoLPAT1-CoLPAT16。",
        44: "根据普通油茶基因组注释文件提取 CoLPAT 基因的染色体编号、起始位点和终止位点等信息，利用 TBtools 与 Python 绘制染色体定位图，以展示 CoLPAT 家族成员在普通油茶基因组中的分布特征。",
        52: "下载拟南芥 TAIR10 基因组注释文件和蛋白序列，提取 AtLPAT 基因位置信息。采用 DIAMOND 对普通油茶与拟南芥最长蛋白序列进行全蛋白同源比对，并利用 WGDI 识别两物种之间的共线性区块。在此基础上，将 CoLPAT 与 AtLPAT 的坐标映射到共线性结果中，筛选涉及 CoLPAT 或 AtLPAT 的相关区块；同时根据最佳同源比对关系绘制 AtLPAT 与 CoLPAT 的同源连接示意图，用于辅助解释系统发育和共线性结果。",
        56: "从 GEO 数据库下载普通油茶公开转录组项目 GSE190644 的 FPKM 表达矩阵及转录本序列。将转录本序列翻译为蛋白序列后，利用 PF01553 隐马尔可夫模型筛选 LPAT 同源转录本，并通过 DIAMOND 将 16 个 CoLPAT 蛋白映射到公共表达数据中的对应转录本。表达量采用 log2(FPKM+1) 转换，并按基因进行行标准化后绘制高变异材料表达热图；同时统计 221 份材料中各 CoLPAT 同源转录本的中位 FPKM，以反映公开群体中的整体表达水平差异。",
        66: "将普通油茶 16 个 CoLPAT 蛋白与拟南芥 AtLPAT 蛋白进行多序列比对并构建系统发育树。结果显示，CoLPAT 成员主要与 AtLPAT1、AtLPAT2 和 AtLPAT4 形成较近的聚类关系，其中 9 个成员与 AtLPAT1 的亲缘关系最近，4 个成员与 AtLPAT2 关系较近，3 个成员与 AtLPAT4 关系较近，说明普通油茶 LPAT 家族在拟南芥已知 LPAT 分支上发生了不同程度的扩张。系统发育结果与后续保守结构域和同源比对分析基本一致，为 CoLPAT 成员的功能推测提供了进化依据。",
        82: "根据 AtLPAT 与 CoLPAT 最佳同源关系构建的连接示意图显示，普通油茶 LPAT 成员主要对应于拟南芥 AtLPAT1、AtLPAT2 和 AtLPAT4 三类分支，其中 AtLPAT1 相关成员最多（图6）。进一步利用 WGDI 对普通油茶与拟南芥全基因组共线性进行分析，共检测到 1370 个种间共线性区块，其中有 5 个区块涉及 CoLPAT 或 AtLPAT 位点（图7）。在严格共线性证据下，CoLPAT16 与 AtLPAT3 位于同一共线性区块；CoLPAT2、CoLPAT4、CoLPAT6 和 CoLPAT7 也落在种间共线性区块中，但其对应拟南芥区段未覆盖 AtLPAT 位点。其余 11 个成员未检测到与 AtLPAT 位点直接对应的共线性证据，表明 CoLPAT 家族与拟南芥 LPAT 之间的保守共线性关系相对有限，最佳同源关系可作为系统发育和功能推断的补充依据，但不能直接等同于严格的种间共线性。",
        94: "基于 GSE190644 公开表达矩阵的补充分析表明，16 个 CoLPAT 成员均可映射到油茶公共转录本中的 LPAT 同源序列，其中 13 个为高置信映射，3 个为低置信映射。表达热图显示，在表达变异较大的 40 份材料中，不同 CoLPAT 同源转录本的表达模式存在明显分化，但同一 AtLPAT 亲缘类别成员之间仍表现出一定相似性（图9）。对 221 份材料的中位 FPKM 统计进一步表明，CoLPAT2、CoLPAT4 和 CoLPAT1 对应转录本整体表达水平较高，CoLPAT8、CoLPAT6 和 CoLPAT7 次之，而 CoLPAT13-CoLPAT16 整体表达较低（图10）。需要说明的是，本研究采用的是公开群体表达数据，因此该结果反映的是材料间差异及同源转录本整体表达趋势，不能直接替代特定组织、特定发育时期或 qRT-PCR 验证结果。",
        98: "LPAT 是植物甘油脂合成途径中的关键酰基转移酶，在磷脂酸形成、膜脂重塑和油脂积累调控中具有重要作用。前人研究表明，不同植物中 LPAT 家族成员数量和功能分化程度存在明显差异，且与基因组倍性、复制历史及组织表达特征密切相关。本研究在普通油茶基因组中系统鉴定出 16 个 CoLPAT 成员，数量高于拟南芥等模式植物，提示多倍化背景下普通油茶 LPAT 家族经历了扩张与保留。",
        99: "系统发育、保守结构域和基因结构分析表明，普通油茶 LPAT 家族整体较为保守，所有成员均保留 PF01553 核心结构域，但在内含子数目、蛋白长度及二级结构组成上仍存在差异，提示其在维持酶学核心功能的同时可能发生了表达调控或底物偏好上的分化。结合最佳同源关系和 WGDI 共线性结果可以看出，普通油茶 LPAT 成员与拟南芥 AtLPAT1、AtLPAT2、AtLPAT4 相关分支的联系更为紧密，而严格共线性证据主要集中于少数位点，说明该家族在不同物种间既保留了共同祖先信息，也经历了较强的重排与分化。",
        100: "公开表达数据进一步显示，CoLPAT 家族成员之间存在明显表达分层，其中 CoLPAT1、CoLPAT2、CoLPAT4、CoLPAT6、CoLPAT7 和 CoLPAT8 的同源转录本在群体中整体表达较高，可作为后续开展组织特异性表达检测、酶活验证和功能互作分析的优先候选对象。不同成员在群体材料中的表达水平差异，提示其在油茶脂质代谢、膜脂更新或发育调控中的作用可能并不完全相同。",
        101: "受限于当前可获得的数据类型，本研究尚未构建普通油茶不同组织或种子发育时期的真实 TPM/FPKM 矩阵，公开表达结果更多反映群体材料间的总体表达趋势。因此，后续仍需结合特定组织样品、种子发育时期表达谱、亚细胞定位及遗传转化实验，对 CoLPAT 候选成员的生物学功能进行进一步验证。",
        104: "本研究基于普通油茶“长林40号”基因组共鉴定出 16 个 CoLPAT 基因家族成员，系统揭示了其染色体定位、系统发育关系、保守结构域、基因结构和蛋白二级结构特征。所有成员均含 PF01553 保守结构域，家族整体较为保守，但不同成员在结构组成上表现出一定差异。结合 GSE190644 公开表达矩阵，16 个 CoLPAT 成员均可映射到 LPAT 同源转录本，其中 CoLPAT1、CoLPAT2、CoLPAT4、CoLPAT6、CoLPAT7 和 CoLPAT8 整体表达相对较高。种间比较分析显示，普通油茶与拟南芥之间存在与 CoLPAT/AtLPAT 相关的共线性区块，且 CoLPAT16 与 AtLPAT3 具有严格共线性关系。上述结果为普通油茶 LPAT 家族的功能研究、脂质代谢调控解析及优异油脂性状分子改良提供了基础资料。",
    }

    for index, text in replacements.items():
        set_paragraph_text(doc.paragraphs[index], text, size=12, first_line=True, center=False)

    set_caption_text(doc.paragraphs[89], "图8  普通油茶 CoLPAT 蛋白二级结构预测")
    set_caption_text(doc.paragraphs[96], "图9  GSE190644 中 CoLPAT 同源转录本在高变异材料中的表达热图")

    delete_from_heading(doc, "公开表达与种间共线性补充分析")

    for index in [106, 105, 102]:
        remove_paragraph(doc.paragraphs[index])

    # Reacquire paragraphs after deletion.
    idx_fig6_caption = find_paragraph_index(doc, "图6")
    idx_fig9_caption = find_paragraph_index(doc, "图9")
    fig6_cap = doc.paragraphs[idx_fig6_caption]
    fig9_cap = doc.paragraphs[idx_fig9_caption]

    fig7_path = FIG_DIR / "fig6_species_collinearity_WGDI_CoLPAT_AtLPAT.png"
    fig10_path = FIG_DIR / "fig9_expression_median_FPKM_GSE190644.png"

    add_picture_with_caption(
        fig6_cap,
        fig7_path,
        "图7  普通油茶 CoLPAT 与拟南芥 AtLPAT 相关 WGDI 共线性块",
        width_cm=15.8,
    )
    add_picture_with_caption(
        fig9_cap,
        fig10_path,
        "图10  CoLPAT 同源转录本在 221 份油茶材料中的中位 FPKM",
        width_cm=14.5,
    )

    out = OUT_DIR / "白垚 毕业论文4(2)_终改插图完善版.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
